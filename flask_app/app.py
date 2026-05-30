"""
IntegriCheck v6 — Flask Web Application
========================================
DROP-IN REPLACEMENT for: flask_app/app.py

FIXES vs v5:
  1. /api/report/<id>  — was calling undefined variables (plag_data, ai_data)
                         now correctly calls both generate_plagiarism_report + generate_ai_report
  2. analyze_ai_detection → renamed to analyze_ai (matches new engine v5)
  3. _build_top_sources   — now reads from new engine's 'top_sources' list directly
                            (engine v6 already returns shaped top_sources, no re-mapping needed)
  4. _build_report_data   — now correctly feeds new engine output keys:
                            'similarity_pct', 'highlights', 'match_groups', 'database_pct',
                            'ai_pct', 'ai_highlights', 'ai_only_count'
  5. Both PDF reports now get ALL required keys (student_name, submission_id, etc.)
  6. /api/report/<id>/plagiarism and /api/report/<id>/ai — return separate PDFs correctly
  7. /api/report/<id>       — returns BOTH PDFs as a zip (combined download)
  8. Error handling improved — no more 500s from missing keys
  9. CORS + file size config kept
"""

import io
import os
import sys
import json
import zipfile
import hashlib
import logging
import tempfile
import traceback
from datetime import datetime

from flask import Flask, request, jsonify, render_template, send_file
from flask_cors import CORS

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# ── App setup ─────────────────────────────────────────────────────────────────
app = Flask(__name__)
CORS(app)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024   # 100 MB

logging.basicConfig(level=logging.INFO,
                    format='[IntegriCheck] %(levelname)s — %(message)s')
logger = logging.getLogger('integricheck')

REPORTS_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'reports')
os.makedirs(REPORTS_DIR, exist_ok=True)


# ══════════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION
# ══════════════════════════════════════════════════════════════════════════════

def _human_size(b: int) -> str:
    for unit in ['B', 'KB', 'MB', 'GB']:
        if b < 1024:
            return f'{b:.1f} {unit}'
        b /= 1024
    return f'{b:.1f} GB'


def _clean(text: str) -> str:
    import re
    if not text:
        return ''
    text = re.sub(r'-\n([a-z])', r'\1', text)
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    raw_paras = re.split(r'\n{2,}', text)
    cleaned = []
    for para in raw_paras:
        lines  = [l.strip() for l in para.split('\n')]
        joined = ' '.join(l for l in lines if l)
        import re as _re
        joined = _re.sub(r'  +', ' ', joined).strip()
        if joined:
            cleaned.append(joined)
    return '\n\n'.join(cleaned)


def _pdf_extract(path: str) -> dict:
    # Try pdfplumber first (better layout)
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            num_pages = len(pdf.pages)
            for page in pdf.pages:
                t = page.extract_text(x_tolerance=2, y_tolerance=3)
                if t and t.strip():
                    pages.append(t.strip())
        if pages:
            return {'text': '\n\n'.join(pages), 'num_pages': num_pages}
    except Exception as e:
        logger.warning(f'pdfplumber failed: {e}')

    # Fallback: pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(path)
        pages  = []
        for page in reader.pages:
            t = page.extract_text()
            if t and t.strip():
                pages.append(t.strip())
        if pages:
            return {'text': '\n\n'.join(pages), 'num_pages': len(reader.pages)}
        return {'text': '', 'num_pages': 0,
                'error': 'PDF appears scanned/image-based — no text found.'}
    except Exception as e:
        return {'text': '', 'num_pages': 0, 'error': f'PDF extraction failed: {e}'}


def _docx_extract(path: str) -> dict:
    try:
        from docx import Document
        doc   = Document(path)
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_t = '  |  '.join(c.text.strip()
                                     for c in row.cells if c.text.strip())
                if row_t:
                    paras.append(row_t)
        text = '\n\n'.join(paras)
        return {'text': text, 'num_pages': max(1, len(paras) // 15)}
    except Exception as e:
        return {'text': '', 'num_pages': 0, 'error': f'DOCX extraction failed: {e}'}


def _txt_extract(path: str) -> dict:
    try:
        with open(path, 'r', encoding='utf-8', errors='replace') as f:
            text = f.read()
        return {'text': text, 'num_pages': max(1, len(text.split()) // 400)}
    except Exception as e:
        return {'text': '', 'num_pages': 0, 'error': f'Text read failed: {e}'}


def _extract_text(file) -> dict:
    """Extract text from uploaded file. Returns dict with text + metadata."""
    fn     = (file.filename or '').lower()
    ext    = os.path.splitext(fn)[-1].lower()
    suffix = ext if ext else '.tmp'

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        file.save(tmp)
        tmp_path = tmp.name

    try:
        size_bytes = os.path.getsize(tmp_path)
        file_size  = _human_size(size_bytes)

        if ext == '.pdf':
            raw = _pdf_extract(tmp_path)
        elif ext in ('.docx', '.doc'):
            raw = _docx_extract(tmp_path)
        elif ext == '.txt':
            raw = _txt_extract(tmp_path)
        else:
            return {
                'text': '', 'word_count': 0, 'char_count': 0,
                'num_pages': 0, 'file_size': file_size,
                'error': f'Unsupported file type: "{ext}". Please upload PDF, DOCX, or TXT.'
            }

        err = raw.get('error')
        if err:
            return {'text': '', 'word_count': 0, 'char_count': 0,
                    'num_pages': 0, 'file_size': file_size, 'error': err}

        text = _clean(raw.get('text', ''))
        return {
            'text':       text,
            'word_count': len(text.split()),
            'char_count': len(text),
            'num_pages':  raw.get('num_pages', 1),
            'file_size':  file_size,
            'error':      None,
        }
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


# ══════════════════════════════════════════════════════════════════════════════
# REPORT DATA BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def _report_id(seed: str = '') -> str:
    """Generate short unique report ID."""
    return hashlib.md5(f'{seed}{datetime.now().isoformat()}'.encode()
                       ).hexdigest()[:8].upper()


def _safe_int(val, default=0) -> int:
    try:
        return int(val)
    except Exception:
        return default


def _fmt_number(val) -> str:
    try:
        return f'{int(val):,}'
    except Exception:
        return str(val)


def _build_plagiarism_report_data(
        plag_result: dict,
        meta: dict,
        text: str,
        report_id: str,
        extracted: dict = None) -> dict:
    """
    Build the exact dict that generate_plagiarism_report() expects.

    Engine v6 already returns:
        similarity_pct, highlights, match_groups, database_pct,
        integrity_flags, top_sources, full_text
    We just add document metadata on top.
    """
    extracted = extracted or {}

    # --- top_sources: new engine returns correctly shaped list ---
    top_sources = plag_result.get('top_sources', [])

    # Re-number ranks 1-based just in case
    for i, s in enumerate(top_sources):
        s['rank'] = i + 1

    word_count = (extracted.get('word_count')
                  or plag_result.get('word_count')
                  or len(text.split()))
    char_count = extracted.get('char_count') or len(text)
    num_pages  = (extracted.get('num_pages')
                  or max(1, len(text) // 3000))

    return {
        # ── Document metadata ──────────────────────────────────────────────
        'doc_title':       meta.get('doc_title', 'Submitted Document'),
        'student_name':    meta.get('student_name', 'Student'),
        'submission_id':   report_id,
        'submission_date': meta.get('submission_date',
                                    datetime.now().strftime('%b %d, %Y, %I:%M %p GMT+5:30')),
        'download_date':   datetime.now().strftime('%b %d, %Y, %I:%M %p GMT+5:30'),
        'file_name':       meta.get('file_name', extracted.get('filename', 'document.pdf')),
        'file_size':       meta.get('file_size', extracted.get('file_size', 'N/A')),
        'page_count':      num_pages,
        'word_count':      word_count,
        'char_count':      char_count,

        # ── Engine results (pass-through from analyze_plagiarism) ──────────
        'similarity_pct':  _safe_int(plag_result.get('similarity_pct', 0)),
        'full_text':       text or plag_result.get('full_text', ''),
        'highlights':      plag_result.get('highlights', []),
        'match_groups':    plag_result.get('match_groups', {
            'not_cited':         {'count': 0, 'pct': 0},
            'missing_quotation': {'count': 0, 'pct': 0},
            'missing_citation':  {'count': 0, 'pct': 0},
            'cited_and_quoted':  {'count': 0, 'pct': 0},
        }),
        'database_pct':    plag_result.get('database_pct', {
            'Internet': 0, 'Publication': 0, 'Student': 0,
        }),
        'integrity_flags': _safe_int(plag_result.get('integrity_flags', 0)),
        'top_sources':     top_sources,
    }


def _build_ai_report_data(
        ai_result: dict,
        meta: dict,
        text: str,
        report_id: str,
        extracted: dict = None) -> dict:
    """
    Build the exact dict that generate_ai_report() expects.

    Engine v5 already returns:
        ai_pct, ai_label, full_text, ai_highlights, ai_only_count,
        sentence_scores, feature_values
    We just add document metadata on top.
    """
    extracted = extracted or {}

    word_count = (extracted.get('word_count')
                  or ai_result.get('word_count')
                  or len(text.split()))
    char_count = extracted.get('char_count') or len(text)
    num_pages  = (extracted.get('num_pages')
                  or max(1, len(text) // 3000))

    return {
        # ── Document metadata ──────────────────────────────────────────────
        'doc_title':       meta.get('doc_title', 'Submitted Document'),
        'student_name':    meta.get('student_name', 'Student'),
        'submission_id':   report_id,
        'submission_date': meta.get('submission_date',
                                    datetime.now().strftime('%b %d, %Y, %I:%M %p GMT+5:30')),
        'download_date':   datetime.now().strftime('%b %d, %Y, %I:%M %p GMT+5:30'),
        'file_name':       meta.get('file_name', extracted.get('filename', 'document.pdf')),
        'file_size':       meta.get('file_size', extracted.get('file_size', 'N/A')),
        'page_count':      num_pages,
        'word_count':      word_count,
        'char_count':      char_count,

        # ── Engine results (pass-through from analyze_ai) ──────────────────
        'ai_pct':          _safe_int(ai_result.get('ai_pct', 0)),
        'ai_label':        ai_result.get('ai_label', 'LOW AI RISK'),
        'full_text':       text or ai_result.get('full_text', ''),
        'ai_highlights':   ai_result.get('ai_highlights', []),
        'ai_only_count':   _safe_int(ai_result.get('ai_only_count', 0)),
        'sentence_scores': ai_result.get('sentence_scores', []),
        'feature_values':  ai_result.get('feature_values', {}),
    }


def _generate_pdf_paths(report_id: str):
    plag_path = os.path.join(REPORTS_DIR, f'Plagiarism_Report_{report_id}.pdf')
    ai_path   = os.path.join(REPORTS_DIR, f'AI_Writing_Report_{report_id}.pdf')
    return plag_path, ai_path


# ══════════════════════════════════════════════════════════════════════════════
# FLASK ROUTES
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/dashboard')
def dashboard():
    return render_template('dashboard.html')


@app.route('/api/health')
def health():
    return jsonify({
        'status':  'ok',
        'service': 'IntegriCheck',
        'version': '6.0.0',
        'time':    datetime.now().isoformat(),
    })


# ── Analyze: plain text ───────────────────────────────────────────────────────

@app.route('/api/analyze/text', methods=['POST'])
def analyze_text():
    """
    POST /api/analyze/text
    Body (JSON): { "text": "...", "meta": { "student_name": "...", ... } }
    Returns: full analysis result JSON
    """
    try:
        data = request.get_json(silent=True) or {}
        text = (data.get('text') or '').strip()

        if not text:
            return jsonify({'error': 'No text provided.'}), 400
        if len(text) < 50:
            return jsonify({'error': 'Text too short — minimum 50 characters required.'}), 400

        from src.plagiarism.engine1   import analyze_plagiarism
        from src.ai_detection.engine1 import analyze_ai

        logger.info(f'Analyzing text ({len(text)} chars)…')
        plag_result = analyze_plagiarism(text)
        ai_result   = analyze_ai(text)
        rid         = _report_id(text[:60])

        # ── Backward-compat aliases for older frontend builds ──────────────
        plag_result['final_score']       = plag_result.get('similarity_pct', 0)
        plag_result['char_highlights']   = plag_result.get('highlights', [])
        ai_result['ai_probability']      = ai_result.get('ai_pct', 0)
        ai_result['sentence_highlights'] = [
            {**s, 'is_ai': s.get('is_ai', False)}
            for s in ai_result.get('sentence_scores', [])
        ]

        return jsonify({
            'report_id':         rid,
            'text_length':       len(text),
            'word_count':        len(text.split()),
            'plagiarism':        plag_result,
            'ai_detection':      ai_result,
            'timestamp':         datetime.now().isoformat(),
        })

    except ImportError as e:
        logger.error(f'Engine import failed: {e}')
        return jsonify({'error': f'Engine not available: {e}'}), 500
    except Exception as e:
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e), 'trace': traceback.format_exc()[-600:]}), 500


# ── Analyze: file upload ──────────────────────────────────────────────────────

@app.route('/api/analyze/file', methods=['POST'])
def analyze_file():
    """
    POST /api/analyze/file
    Form-data: file=<upload>, meta=<JSON string> (optional)
    Returns: full analysis result JSON
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded. Use key "file".'}), 400

        f = request.files['file']
        if not f or not f.filename:
            return jsonify({'error': 'No file selected.'}), 400

        # Optional metadata from form field
        meta_raw = request.form.get('meta', '{}')
        try:
            meta = json.loads(meta_raw)
        except Exception:
            meta = {}

        meta.setdefault('file_name', f.filename)

        logger.info(f'Extracting text from: {f.filename}')
        extracted = _extract_text(f)

        if extracted.get('error'):
            return jsonify({'error': extracted['error']}), 400

        text = extracted['text']
        if len(text.strip()) < 50:
            return jsonify({'error': 'Could not extract enough text from the file (min 50 chars).'}), 400

        from src.plagiarism.engine1   import analyze_plagiarism
        from src.ai_detection.engine1 import analyze_ai

        logger.info(f'Running analysis on {extracted["word_count"]} words…')
        plag_result = analyze_plagiarism(text)
        ai_result   = analyze_ai(text)
        rid         = _report_id(f.filename + text[:40])

        logger.info(f'PLAG RESULT = similarity_pct:{plag_result.get("similarity_pct")}, '
                    f'highlights:{len(plag_result.get("highlights",[]))}, '
                    f'top_sources:{len(plag_result.get("top_sources",[]))}')
        logger.info(f'AI RESULT = ai_pct:{ai_result.get("ai_pct")}, '
                    f'ai_highlights:{len(ai_result.get("ai_highlights",[]))}')

        # ── Backward-compat aliases for older frontend builds ──────────────
        plag_result['final_score']       = plag_result.get('similarity_pct', 0)
        plag_result['char_highlights']   = plag_result.get('highlights', [])
        ai_result['ai_probability']      = ai_result.get('ai_pct', 0)
        ai_result['sentence_highlights'] = [
            {**s, 'is_ai': s.get('is_ai', False)}
            for s in ai_result.get('sentence_scores', [])
        ]

        return jsonify({
            'report_id':    rid,
            'filename':     f.filename,
            'word_count':   extracted['word_count'],
            'char_count':   extracted['char_count'],
            'num_pages':    extracted['num_pages'],
            'file_size':    extracted['file_size'],
            'text_length':  len(text),          # total chars extracted
            'plagiarism':   plag_result,
            'ai_detection': ai_result,
            'timestamp':    datetime.now().isoformat(),
        })

    except ImportError as e:
        logger.error(f'Engine import failed: {e}')
        return jsonify({'error': f'Engine not available: {e}'}), 500
    except Exception as e:
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e), 'trace': traceback.format_exc()[-600:]}), 500


# ── Download: Plagiarism PDF ──────────────────────────────────────────────────

@app.route('/api/report/<report_id>/plagiarism', methods=['POST'])
def download_plagiarism_report(report_id):
    """
    POST /api/report/<id>/plagiarism
    Body (JSON): {
        "text": "...",
        "plagiarism": { ...analyze_plagiarism() result... },
        "meta": { "student_name": "...", "doc_title": "..." },
        "word_count": int, "char_count": int, "num_pages": int, "file_size": "..."
    }
    Returns: PDF file download
    """
    try:
        from src.utils.report_generator import generate_plagiarism_report

        body        = request.get_json(silent=True) or {}
        text        = body.get('text', '')
        plag_result = body.get('plagiarism', {})
        meta        = body.get('meta', {})
        extracted   = {
            'word_count': body.get('word_count'),
            'char_count': body.get('char_count'),
            'num_pages':  body.get('num_pages'),
            'file_size':  body.get('file_size'),
            'filename':   body.get('filename', ''),
        }

        report_data = _build_plagiarism_report_data(
            plag_result, meta, text, report_id, extracted)

        plag_path, _ = _generate_pdf_paths(report_id)
        generate_plagiarism_report(report_data, plag_path)

        logger.info(f'Plagiarism PDF generated: {plag_path}')
        return send_file(
            plag_path,
            as_attachment=True,
            download_name=f'Plagiarism_Report_{report_id}.pdf',
            mimetype='application/pdf',
        )

    except Exception as e:
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e), 'trace': traceback.format_exc()[-800:]}), 500


# ── Download: AI Detection PDF ────────────────────────────────────────────────

@app.route('/api/report/<report_id>/ai', methods=['POST'])
def download_ai_report(report_id):
    """
    POST /api/report/<id>/ai
    Body (JSON): {
        "text": "...",
        "ai_detection": { ...analyze_ai() result... },
        "meta": { "student_name": "...", "doc_title": "..." },
        "word_count": int, "char_count": int, "num_pages": int, "file_size": "..."
    }
    Returns: PDF file download
    """
    try:
        from src.utils.report_generator import generate_ai_report

        body      = request.get_json(silent=True) or {}
        text      = body.get('text', '')
        ai_result = body.get('ai_detection', {})
        meta      = body.get('meta', {})
        extracted = {
            'word_count': body.get('word_count'),
            'char_count': body.get('char_count'),
            'num_pages':  body.get('num_pages'),
            'file_size':  body.get('file_size'),
            'filename':   body.get('filename', ''),
        }

        report_data = _build_ai_report_data(
            ai_result, meta, text, report_id, extracted)

        _, ai_path = _generate_pdf_paths(report_id)
        generate_ai_report(report_data, ai_path)

        logger.info(f'AI PDF generated: {ai_path}')
        return send_file(
            ai_path,
            as_attachment=True,
            download_name=f'AI_Writing_Report_{report_id}.pdf',
            mimetype='application/pdf',
        )

    except Exception as e:
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e), 'trace': traceback.format_exc()[-800:]}), 500


# ── Download: Both PDFs as ZIP ────────────────────────────────────────────────

@app.route('/api/report/<report_id>', methods=['POST'])
def download_both_reports(report_id):
    """
    POST /api/report/<id>
    Body (JSON): {
        "text": "...",
        "plagiarism":   { ...analyze_plagiarism() result... },
        "ai_detection": { ...analyze_ai() result... },
        "meta": { "student_name": "...", "doc_title": "..." },
        "word_count": int, "char_count": int, "num_pages": int, "file_size": "..."
    }
    Returns: ZIP containing both Plagiarism PDF + AI Writing PDF
    """
    try:
        from src.utils.report_generator import generate_plagiarism_report, generate_ai_report

        body        = request.get_json(silent=True) or {}
        text        = body.get('text', '')
        plag_result = body.get('plagiarism', {})
        ai_result   = body.get('ai_detection', {})
        meta        = body.get('meta', {})
        extracted   = {
            'word_count': body.get('word_count'),
            'char_count': body.get('char_count'),
            'num_pages':  body.get('num_pages'),
            'file_size':  body.get('file_size'),
            'filename':   body.get('filename', ''),
        }

        # Build both report data dicts
        plag_data = _build_plagiarism_report_data(
            plag_result, meta, text, report_id, extracted)
        ai_data   = _build_ai_report_data(
            ai_result, meta, text, report_id, extracted)

        # Generate both PDFs
        plag_path, ai_path = _generate_pdf_paths(report_id)
        generate_plagiarism_report(plag_data, plag_path)
        generate_ai_report(ai_data, ai_path)

        logger.info(f'Both PDFs generated for report {report_id}')

        # Bundle into ZIP in memory
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.write(plag_path, arcname=f'Plagiarism_Report_{report_id}.pdf')
            zf.write(ai_path,   arcname=f'AI_Writing_Report_{report_id}.pdf')
        zip_buffer.seek(0)

        return send_file(
            zip_buffer,
            as_attachment=True,
            download_name=f'IntegriCheck_Reports_{report_id}.zip',
            mimetype='application/zip',
        )

    except Exception as e:
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e), 'trace': traceback.format_exc()[-800:]}), 500


# ── Bonus: Re-analyze stored text and download reports in one shot ────────────

@app.route('/api/analyze-and-download', methods=['POST'])
def analyze_and_download():
    """
    POST /api/analyze-and-download
    Form-data OR JSON:
        file   = uploaded file  (optional — if text not provided)
        text   = raw text       (optional — if file not provided)
        meta   = JSON string    { student_name, doc_title }
        report = 'plagiarism' | 'ai' | 'both' (default: 'both')

    Runs full analysis and immediately returns PDF(s).
    Frontend can call this single endpoint to get reports without a separate analyze step.
    """
    try:
        from src.plagiarism.engine1   import analyze_plagiarism
        from src.ai_detection.engine1 import analyze_ai
        from src.utils.report_generator import generate_plagiarism_report, generate_ai_report

        # --- Get text ---
        text      = ''
        extracted = {}
        if 'file' in request.files:
            f = request.files['file']
            if f and f.filename:
                extracted = _extract_text(f)
                if extracted.get('error'):
                    return jsonify({'error': extracted['error']}), 400
                text = extracted.get('text', '')
                extracted['filename'] = f.filename

        if not text:
            data = request.get_json(silent=True) or {}
            text = (data.get('text') or request.form.get('text') or '').strip()

        if len(text.strip()) < 50:
            return jsonify({'error': 'Not enough text to analyze (min 50 chars).'}), 400

        # --- Metadata ---
        meta_raw = (request.form.get('meta') or
                    (request.get_json(silent=True) or {}).get('meta', '{}'))
        if isinstance(meta_raw, str):
            try:
                meta = json.loads(meta_raw)
            except Exception:
                meta = {}
        else:
            meta = meta_raw or {}
        meta.setdefault('file_name', extracted.get('filename', 'document'))

        # --- Report type ---
        report_type = (request.form.get('report') or
                       (request.get_json(silent=True) or {}).get('report', 'both'))

        # --- Run analysis ---
        rid = _report_id(text[:60])
        logger.info(f'analyze-and-download: report_id={rid}, type={report_type}')

        plag_result = analyze_plagiarism(text) if report_type in ('plagiarism', 'both') else {}
        ai_result   = analyze_ai(text)         if report_type in ('ai', 'both')         else {}

        plag_path, ai_path = _generate_pdf_paths(rid)

        if report_type == 'plagiarism':
            plag_data = _build_plagiarism_report_data(plag_result, meta, text, rid, extracted)
            generate_plagiarism_report(plag_data, plag_path)
            return send_file(plag_path, as_attachment=True,
                             download_name=f'Plagiarism_Report_{rid}.pdf',
                             mimetype='application/pdf')

        if report_type == 'ai':
            ai_data = _build_ai_report_data(ai_result, meta, text, rid, extracted)
            generate_ai_report(ai_data, ai_path)
            return send_file(ai_path, as_attachment=True,
                             download_name=f'AI_Writing_Report_{rid}.pdf',
                             mimetype='application/pdf')

        # Both — return ZIP
        plag_data = _build_plagiarism_report_data(plag_result, meta, text, rid, extracted)
        ai_data_d = _build_ai_report_data(ai_result, meta, text, rid, extracted)
        generate_plagiarism_report(plag_data, plag_path)
        generate_ai_report(ai_data_d, ai_path)

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            zf.write(plag_path, arcname=f'Plagiarism_Report_{rid}.pdf')
            zf.write(ai_path,   arcname=f'AI_Writing_Report_{rid}.pdf')
        zip_buffer.seek(0)

        return send_file(zip_buffer, as_attachment=True,
                         download_name=f'IntegriCheck_Reports_{rid}.zip',
                         mimetype='application/zip')

    except ImportError as e:
        logger.error(f'Engine import failed: {e}')
        return jsonify({'error': f'Engine not available: {e}'}), 500
    except Exception as e:
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e), 'trace': traceback.format_exc()[-800:]}), 500


# ── 404 / 405 handlers ────────────────────────────────────────────────────────

@app.errorhandler(404)
def not_found(e):
    return jsonify({'error': 'Endpoint not found.', 'hint': 'Check /api/health'}), 404


@app.errorhandler(405)
def method_not_allowed(e):
    return jsonify({'error': 'Method not allowed on this endpoint.'}), 405


@app.errorhandler(413)
def file_too_large(e):
    return jsonify({'error': 'File too large. Maximum size is 100 MB.'}), 413


# ══════════════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=False)
